from __future__ import annotations

import argparse
import hashlib
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import cobra
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_RUN = ROOT / "runs" / "10h2da-three-system-20260717"
DEFAULT_OUTPUT = ROOT / "07_reports"

SYSTEMS = {
    "yeast_gem": {
        "label": "yeast-GEM",
        "filename": "10H2DA_yeast-GEM_中文报告.docx",
        "run_id": "10h2da-yeast_gem",
        "conclusion": "yeast-GEM 的两条路线目标通量均为 0.243015，禁氧结果为 0；它提供保守基线，但终端两步仍是候选扩展。",
    },
    "yeast_metatwin": {
        "label": "Yeast-MetaTwin",
        "filename": "10H2DA_Yeast-MetaTwin_中文报告.docx",
        "run_id": "10h2da-yeast_metatwin",
        "conclusion": "Yeast-MetaTwin 的两条路线目标通量均为 0.274011，约比 yeast-GEM 高 12.75%；禁氧仍有正通量并触发 oxygen bypass 警告，不能解释为真实厌氧生产。",
    },
    "integrated": {
        "label": "Integrated Deployment",
        "filename": "10H2DA_Integrated_Deployment_中文报告.docx",
        "run_id": "10h2da-integrated",
        "conclusion": "Integrated Deployment 使用 Yeast-MetaTwin 底模，因此 FBA 数值同为 0.274011；新增价值是可追溯证据、动力学门禁、工程规则和 hypothetical 构建层，而非第三套计量网络。",
    },
}

REACTION_INTERPRETATION = {
    "r_0399": ("FAA2 / YER015W", "参与", "癸酸 CoA 活化，支持 C10 前体进入 CoA 池；不等于必须改造。"),
    "r_0120": ("POX1 / YGL205W", "参与", "形成 trans-dec-2-enoyl-CoA，是模型内直接前体步骤。"),
    "r_2248": ("FOX2 / YKR009C", "竞争消耗", "水合 trans-dec-2-enoyl-CoA，可能与终端释放竞争；仅为节点，不自动推出敲除。"),
    "r_2295": ("ECI1 / YLR284C", "替代连接", "从另一不饱和 C10-CoA 节点连接到 trans-dec-2-enoyl-CoA；不是终端反应。"),
}

REQUIRED_KEYWORDS = [
    "C10H18O3", "C10H17O3", "S. cerevisiae", "native growth", "FBA", "pFBA", "FVA",
    "r_0399", "r_0120", "r_2248", "r_2295", "TES1", "B8QHP1", "Q9Y8G7",
    "exact evidence count=0", "可复现代码", "workflow_cli.py", "executors.py",
]


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def relative(path: Path) -> str:
    try:
        return path.resolve().relative_to(ROOT.resolve()).as_posix()
    except ValueError:
        return str(path.resolve())


def resolve_model(config: dict[str, Any]) -> Path:
    configured = Path(config["fba_validation"]["model_path"])
    return configured if configured.is_absolute() else (ROOT / configured).resolve()


def collect_reactions(model_path: Path, include_rxn1937: bool) -> list[dict[str, str]]:
    model = cobra.io.load_yaml_model(model_path)
    ids = ["r_0399", "r_0120", "r_2248", "r_2295"] + (["rxn1937"] if include_rxn1937 else [])
    rows = []
    for reaction_id in ids:
        reaction = model.reactions.get_by_id(reaction_id)
        if reaction_id == "rxn1937":
            declared, role, note = "模型动态读取的多基因 OR 规则", "替代连接", "地下反应提供另一条到 s_1507 的连接；预测来源，不能视作 14 个基因都必需。"
        else:
            declared, role, note = REACTION_INTERPRETATION[reaction_id]
        rows.append({
            "id": reaction.id,
            "declared": declared,
            "equation": reaction.reaction,
            "gpr": reaction.gene_reaction_rule or "无 GPR",
            "genes": ", ".join(sorted(gene.id for gene in reaction.genes)) or "无",
            "role": role,
            "note": note,
        })
    return rows


def load_system(run_dir: Path, key: str) -> dict[str, Any]:
    info = SYSTEMS[key]
    config_path = run_dir / "configs" / f"{key}.json"
    artifact_root = run_dir / "runs" / info["run_id"] / "artifacts"
    paths = {
        "config": config_path,
        "fba": artifact_root / "model_feasibility" / "fba_results.json",
        "evidence": artifact_root / "external_evidence" / "external_evidence.json",
        "kinetic": artifact_root / "kinetic_prediction" / "kinetic_predictions.json",
        "engineering": artifact_root / "engineering_feasibility" / "engineering_feasibility.json",
        "construct": artifact_root / "construct_design" / "construct_designs.json",
    }
    missing = [str(path) for path in paths.values() if not path.is_file()]
    if missing:
        raise FileNotFoundError("missing required comparison artifacts: " + ", ".join(missing))
    config = read_json(config_path)
    model_path = resolve_model(config)
    if not model_path.is_file():
        raise FileNotFoundError(f"model not found: {model_path}")
    return {
        "key": key,
        "info": info,
        "paths": {**paths, "model": model_path},
        "config": config,
        "fba": read_json(paths["fba"]),
        "evidence": read_json(paths["evidence"]),
        "kinetic": read_json(paths["kinetic"]),
        "engineering": read_json(paths["engineering"]),
        "construct": read_json(paths["construct"]),
        "reactions": collect_reactions(model_path, key != "yeast_gem"),
    }


def set_font(run, name: str = "Microsoft YaHei", size: float = 9.5, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, value: Any, header: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(value))
    set_font(run, size=8.0, bold=header, color="FFFFFF" if header else "25313C")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        set_cell_shading(cell, "1F5A6A")


def add_table(document: Document, headers: list[str], rows: list[list[Any]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, header=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value)
            if row_index % 2:
                set_cell_shading(cells[index], "EAF1F3")
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, end])


def configure_document(document: Document, label: str, source_run: str) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = False

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color in (("Title", 22, "123B46"), ("Heading 1", 15, "123B46"), ("Heading 2", 12, "B54B2A")):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"10H2DA | {label} | 可追溯计算报告")
    set_font(run, size=8, color="55717A")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"来源 run: {source_run}   |   第 ")
    set_font(run, size=8, color="55717A")
    add_field(footer, "PAGE")
    run = footer.add_run(" 页")
    set_font(run, size=8, color="55717A")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_code(document: Document, code: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F5")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(code.splitlines()):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_font(run, name="Consolas", size=7.7, color="20282D")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def fmt(value: Any, digits: int = 6) -> str:
    if value is None:
        return "N/A"
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def sensitivity(route: dict[str, Any], scenario: str) -> dict[str, Any]:
    return next(item for item in route["sensitivity"] if item["scenario"] == scenario)


def route_rows(fba: dict[str, Any]) -> list[list[str]]:
    rows = []
    names = {"free_acid_route": "游离酸优先路线", "coa_bound_route": "CoA-bound 高风险路线"}
    for route in fba["routes"]:
        target_fva = route["fva"]["reactions"][route["objective_reaction_id"]]
        rows.append([
            names[route["route_id"]],
            f"{route['status']} / {fmt(route['objective_value'])}",
            f"{route['pfba']['status']} / 目标 {fmt(route['pfba']['objective_value'])}\n总绝对通量 {fmt(route['pfba']['total_absolute_flux'])}",
            f"90% optimum\n[{fmt(target_fva['minimum'])}, {fmt(target_fva['maximum'])}]",
            fmt(route["carbon_yield"]["target_carbon_mol_per_substrate_carbon_mol"]),
            f"{sensitivity(route, 'no_carbon')['status']} / {fmt(sensitivity(route, 'no_carbon')['objective_value'])}",
            f"{sensitivity(route, 'anaerobic')['status']} / {fmt(sensitivity(route, 'anaerobic')['objective_value'])}\nbypass={fmt(sensitivity(route, 'anaerobic')['oxygen_bypass_suspected'])}",
            f"flux={fmt(route['cycle_detection']['target_flux'])}\nsuspected={fmt(route['cycle_detection']['cycle_suspected'])}",
        ])
    return rows


def terminal_rows(config: dict[str, Any]) -> list[list[str]]:
    reactions = {item["reaction_id"]: item for item in config["candidate_reactions"]}
    details = {
        "CAND_T2DEC_THIOESTERASE_P": ("游离酸路线 1", "acyl-CoA thioesterase", "H2O", "TES1/PTE1 / YJR019C / P41903", "S. cerevisiae", "C:近底物/酶家族", "未验证 trans-2-enoyl-CoA 特异性；区室与前体可达性。"),
        "CAND_T2DEC_OMEGA_HYDROXYLASE_P": ("游离酸路线 2", "fatty-acid omega-hydroxylase", "NADPH、O2、heme；CYP52 需 CPR", "cyp52M1 / B8QHP1 + NCP1/YHR042W/P16603 或兼容 CPR；替代 CYP505/Q9Y8G7 self-sufficient", "Starmerella bombicola；CPR 为 S. cerevisiae；CYP505 为 Fusarium oxysporum", "D:酶家族候选", "无 10H2DA exact evidence；P450 底物特异性、电子传递、膜定位。"),
        "CAND_T2DEC_COA_OMEGA_HYDROXYLASE_P": ("CoA-bound 路线 1", "CoA-bound omega-hydroxylase", "NADPH、O2、heme；CYP52 需 CPR", "cyp52M1 / B8QHP1 + NCP1/YHR042W/P16603 或兼容 CPR；替代 CYP505/Q9Y8G7 self-sufficient", "同上", "D:酶家族候选", "最高风险：现有 omega-hydroxylase 证据偏游离脂肪酸，CoA-bound 底物支持更弱。"),
        "CAND_10H2DA_COA_THIOESTERASE_P": ("CoA-bound 路线 2", "acyl-CoA thioesterase", "H2O", "TES1/PTE1 / YJR019C / P41903", "S. cerevisiae", "D:酶家族候选", "未验证羟基化 enoyl-CoA 特异性；区室与产物释放。"),
    }
    rows = []
    for route in config["routes"]:
        for reaction_id in route["reaction_ids"]:
            reaction = reactions[reaction_id]
            step, enzyme, cofactor, candidate, source, evidence, risk = details[reaction_id]
            rows.append([step, reaction_id, reaction["equation"], enzyme, ", ".join(reaction["enzyme_ec_numbers"]), cofactor, candidate, source, evidence, risk])
    return rows


def add_cover(document: Document, data: dict[str, Any], run_dir: Path) -> None:
    label = data["info"]["label"]
    title = document.add_heading("10H2DA 代谢通路预测报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(label)
    set_font(run, size=16, bold=True, color="B54B2A")
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("可复现计算结果、候选酶分层与证据边界")
    set_font(run, size=10.5, color="55717A")
    document.add_paragraph()
    add_table(document, ["项目", "定义"], [
        ["目标", "10-Hydroxy-trans-2-decenoic acid (10H2DA)"],
        ["中性分子", "C10H18O3"],
        ["模型物种", "C10H17O3(-1)，单去质子化羧酸"],
        ["宿主", "S. cerevisiae (Saccharomyces cerevisiae)"],
        ["现有比较 run", relative(run_dir)],
        ["底层模型", f"{data['fba']['model']} | SHA256 {sha256(data['paths']['model'])}"],
        ["报告生成时间", datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")],
    ], [3.3, 13.5])
    paragraph = document.add_paragraph("判读边界：所有终端酶均为候选；exact evidence count=0。FBA 可行性不构成序列、酶活、产量或湿实验验证。")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        set_font(run, size=9, bold=True, color="B54B2A")


def add_results(document: Document, data: dict[str, Any]) -> None:
    fba = data["fba"]
    add_heading(document, "1. 目标与模型范围")
    document.add_paragraph(
        "目标定义为中性 10H2DA C10H18O3；按当前模型质子约定使用 C10H17O3(-1)。宿主为 S. cerevisiae。"
        "两条终端路线是在内存中加入的元素/电荷平衡候选反应，原模型并不原生包含 10H2DA。"
    )
    add_heading(document, "2. 当前实际运行结果")
    add_table(document, ["底层模型", "native growth 状态", "native growth max", "biomass reaction", "10% growth floor"], [[
        fba["model"], fba["native_growth_status"], fmt(fba["native_growth_max"]), fba["biomass_reaction_id"], fmt(fba["biomass_floor"])
    ]])
    add_table(document, ["路线", "FBA 状态/目标", "pFBA", "FVA 目标范围", "碳收率", "无碳源", "禁氧", "闭边界循环"], route_rows(fba))
    document.add_paragraph(
        "碳收率单位为 target carbon mol / substrate carbon mol。无碳源测试均为 0。循环检查为关闭除目标需求外全部边界后的目标通量测试。"
        "oxygen bypass 表示关闭配置的氧交换仍有目标通量，提示模型内部产氧或替代氧连接，不是厌氧生物学证据。"
    )


def add_model_reactions(document: Document, data: dict[str, Any]) -> None:
    add_heading(document, "3. 模型直接前体反应、GPR 与基因")
    add_table(document, ["反应", "模型方程", "实际 GPR（动态读取）", "模型基因", "分类", "解释"], [
        [row["id"], row["equation"], row["gpr"], row["genes"], row["role"], row["note"]] for row in data["reactions"]
    ])
    document.add_paragraph(
        "这里的“参与”表示模型网络可用步骤，“竞争消耗”表示前体分流节点，“替代连接”表示另一条网络连接。"
        "这些标签不把所有 GPR 基因提升为必须改造基因；候选终端反应本次无 GPR，单基因删除分析因此为 not_applicable。"
    )


def add_routes(document: Document, data: dict[str, Any]) -> None:
    add_heading(document, "4. 两条候选终端路径")
    document.add_paragraph(
        "优先实现概念：TES1/PTE1/YJR019C/P41903 负责候选硫酯水解，cyp52M1/B8QHP1 (Starmerella bombicola) 负责候选 omega 羟化，"
        "并配 NCP1/YHR042W/P16603 或经兼容性确认的 CPR。替代实现候选为 Fusarium oxysporum CYP505/Q9Y8G7 self-sufficient 融合酶。"
        "上述全部是候选，exact evidence count=0。"
    )
    add_table(document, ["步骤", "反应 ID", "计量方程", "所需酶", "EC", "辅因子", "候选基因/UniProt", "来源", "证据等级", "主要风险"], terminal_rows(data["config"]))
    document.add_paragraph(
        "路线排序：优先游离酸路线，即先释放 trans-2-decenoate 再 omega 羟化；CoA-bound omega hydroxylation 风险更高，"
        "因为家族证据主要支持游离脂肪酸类底物，不能由相同 FBA 通量推断相同酶学可行性。"
    )


def add_modules(document: Document) -> None:
    add_heading(document, "5. 酵母异源生产所需模块")
    rows = [
        ["C10 前体供给", "提高癸酸/癸酰-CoA 可用性；ACC1、FAS1、FAS2 可作为前体供给工程候选，但不是本次终端 FBA 直接证明的必要基因。", "碳流与链长分布"],
        ["CoA 活化", "FAA2/r_0399 及其他适配 acyl-CoA ligase 候选。", "底物谱、ATP/CoA 负担、区室"],
        ["trans-2-enoyl-CoA 形成", "POX1/r_0120；ECI1/r_2295 是替代连接候选。", "beta-oxidation 继续消耗"],
        ["thioesterase", "TES1/PTE1/P41903 候选，按路线作用于 trans-2-enoyl-CoA 或羟基化 CoA。", "exact 底物活性为 0 条"],
        ["omega-hydroxylase", "cyp52M1/B8QHP1 优先候选；CYP505/Q9Y8G7 self-sufficient 替代候选。", "位置选择性与底物特异性"],
        ["NADPH/O2/heme/CPR", "保障 P450 还原力、氧、heme 与电子传递；NCP1/P16603 或兼容 CPR。", "氧旁路、电子耦联与氧化压力"],
        ["区室/转运", "协调过氧化物酶体、ER、胞质之间的 CoA 酯、游离酸和 P450 可达性。", "跨膜与 CoA 物种不可自由扩散"],
        ["产物输出与耐受", "评估 10H2DA 输出、胞内积累与宿主耐受的独立模块。", "膜与酸胁迫；无本次实验证据"],
        ["分析验证", "区分异构体、游离酸/CoA 形式并核对质量与碳平衡；建立空白和过程对照。", "不在本报告提供湿实验参数"],
    ]
    add_table(document, ["模块", "候选功能/元件", "判读重点"], rows)
    document.add_paragraph("本清单是设计需求与验证边界，不包含培养、转化、表达、取样或分析的湿实验操作参数。")


def add_gene_tiers(document: Document) -> None:
    add_heading(document, "6. 基因分层与优先级")
    add_table(document, ["层级", "候选", "定位与边界"], [
        ["模型直接基因", "FAA2/YER015W；POX1/YGL205W；FOX2/YKR009C；ECI1/YLR284C；MetaTwin 的 rxn1937 ORF 集合", "模型 GPR 直接提取；包含参与、竞争消耗和替代连接，不等于全部必须改造。"],
        ["宿主终端候选", "TES1/PTE1/YJR019C/P41903", "S. cerevisiae acyl-CoA thioesterase 家族候选；exact 终端底物证据为 0。"],
        ["异源核心候选", "cyp52M1/B8QHP1；替代 CYP505/Q9Y8G7", "分别为 Starmerella bombicola CYP52 与 Fusarium oxysporum self-sufficient P450 候选；均未验证 10H2DA exact 反应。"],
        ["辅因子/电子传递候选", "NCP1/YHR042W/P16603 或兼容 CPR；NADPH、O2、heme 供给", "支持 P450，不是终端羟化酶本身；CYP505 自带还原酶域但仍需宿主兼容性。"],
        ["竞争节点", "FOX2/YKR009C 等 beta-oxidation 下游；前体/CoA/还原力竞争", "仅为候选调控节点，需先验证代谢与生长影响。"],
        ["低特异性内源候选", "ERG11、ERG5、DIT2 等", "仅因 P450/氧化酶类别或关键词相关；底物类别距离远，不建议当作已验证 10H2DA 羟化酶。"],
        ["前体供给工程候选", "ACC1、FAS1、FAS2", "可用于 C10 前体供给假设；不是本次终端 FBA 直接证明的必要基因。"],
    ])


def add_shared_layers(document: Document, data: dict[str, Any]) -> None:
    add_heading(document, "7. 证据、动力学与工程层")
    evidence = data["evidence"]
    kinetic = data["kinetic"]
    rules = data["engineering"]["rule_results"]
    construct = data["construct"]
    add_table(document, ["层", "本次实际结果", "可声明范围"], [
        ["external evidence", f"{evidence['record_count']} records / {evidence['exact_match_count']} exact；状态 {evidence['status']}", "记录是候选家族/化合物上下文，不证明 exact 终端反应。"],
        ["kinetic", f"{kinetic['result_counts']['ready']} ready / {kinetic['result_counts']['unsupported']} unsupported；原因：缺序列", "没有可用动力学预测，不伪造序列。"],
        ["engineering rules", "; ".join(f"{row['rule_id']} -> {row['result']} (matched={row['matched']})" for row in rules), "仅用于可追溯优先级；experimentally_validated=false。"],
        ["construct", f"状态 {construct['status']}；{len(construct['designs'])} designs；claims.hypothetical={construct['claims']['hypothetical']}", "construct hypothetical、requires_review，非湿实验验证。"],
    ])
    if data["key"] == "integrated":
        document.add_paragraph(
            "Integrated 明确结果：evidence 126 records/0 exact；kinetic 0 ready/8 unsupported（缺序列）；"
            "工程规则分别触发 requires_manual_evidence_review 与 experimental_validation_not_established；construct hypothetical。"
        )
    else:
        document.add_paragraph(
            "以上 evidence、kinetic、engineering 和 construct 结果来自本次三系统统一调用的共享分析层，不是该底层 GEM 本身携带的能力或原生证据。"
            "本报告仅把共享层作为候选解释上下文，不能据此提高底层模型的酶学证明等级。"
        )


def rerun_command(data: dict[str, Any], run_dir: Path) -> str:
    config = data["paths"]["config"].resolve()
    new_id = f"10h2da-{data['key']}-report-rerun-example"
    return (
        f'python "{(ROOT / "10_generic_target_workflow" / "runtime" / "workflow_cli.py").resolve()}" run '
        f'--config "{config}" --project-root "{ROOT.resolve()}" '
        f'--run-id "{new_id}" --runs-dir "{(run_dir / "report-reruns").resolve()}" '
        f'--output-dir "{(run_dir / "report-rerun-workspaces" / data["key"]).resolve()}"'
    )


def add_reproducibility(document: Document, data: dict[str, Any], run_dir: Path) -> None:
    add_heading(document, "8. 可复现代码")
    document.add_paragraph("CLI 重跑命令（使用当前比较 run 已生成的 config；示例 run-id 为新 ID，避免覆盖现有 run）：")
    add_code(document, rerun_command(data, run_dir))
    document.add_paragraph("模型查询 Python 代码（动态读取模型反应、方程、GPR 与基因）：")
    model_code = f'''import json
from pathlib import Path
import cobra

root = Path(r"{ROOT.resolve()}")
config_path = Path(r"{data['paths']['config'].resolve()}")
config = json.loads(config_path.read_text(encoding="utf-8"))
model_path = (root / config["fba_validation"]["model_path"]).resolve()
model = cobra.io.load_yaml_model(model_path)
reaction_ids = {repr([row['id'] for row in data['reactions']])}
for reaction_id in reaction_ids:
    reaction = model.reactions.get_by_id(reaction_id)
    print(reaction.id, reaction.reaction)
    print("GPR:", reaction.gene_reaction_rule)
    print("genes:", sorted(gene.id for gene in reaction.genes))'''
    add_code(document, model_code)
    document.add_paragraph("结果读取 Python 代码（读取 native growth、两路线 FBA/pFBA/FVA、收率和敏感性）：")
    result_code = f'''import json
from pathlib import Path

path = Path(r"{data['paths']['fba'].resolve()}")
result = json.loads(path.read_text(encoding="utf-8"))
print("native growth:", result["native_growth_status"], result["native_growth_max"])
for route in result["routes"]:
    demand = route["objective_reaction_id"]
    print(route["route_id"], "FBA", route["objective_value"])
    print("pFBA", route["pfba"])
    print("FVA", route["fva"]["reactions"][demand])
    print("carbon yield", route["carbon_yield"])
    print("sensitivity", route["sensitivity"])
    print("cycle", route["cycle_detection"])'''
    add_code(document, result_code)
    document.add_paragraph("实际源代码路径：")
    add_table(document, ["文件", "路径"], [
        ["run_10h2da_three_system_comparison.py", relative(ROOT / "08_runtime" / "run_10h2da_three_system_comparison.py")],
        ["workflow_cli.py", relative(ROOT / "10_generic_target_workflow" / "runtime" / "workflow_cli.py")],
        ["executors.py", relative(ROOT / "10_generic_target_workflow" / "runtime" / "executors.py")],
        ["本报告生成脚本", relative(Path(__file__))],
    ])


def add_conclusion(document: Document, data: dict[str, Any]) -> None:
    add_heading(document, "9. 结论与限制")
    document.add_paragraph(data["info"]["conclusion"])
    document.add_paragraph(
        "两路线 FBA 通量相同只说明当前计量约束下可行，不支持按通量宣称 enzyme specificity。"
        "游离酸路线因 CYP52 家族底物语境更接近而优先；CoA-bound omega hydroxylation 风险更高。"
        "所有终端候选 exact evidence count=0；没有序列输入，没有动力学 ready 结果，也没有实验验证。"
    )
    add_heading(document, "10. 输入与追溯", level=1)
    add_table(document, ["输入", "相对路径", "SHA256"], [[name, relative(path), sha256(path)] for name, path in data["paths"].items()])


def build_report(data: dict[str, Any], run_dir: Path, output_path: Path) -> None:
    document = Document()
    configure_document(document, data["info"]["label"], run_dir.name)
    add_cover(document, data, run_dir)
    document.add_section(WD_SECTION.NEW_PAGE)
    add_results(document, data)
    add_model_reactions(document, data)
    add_routes(document, data)
    add_modules(document)
    add_gene_tiers(document)
    add_shared_layers(document, data)
    add_reproducibility(document, data, run_dir)
    add_conclusion(document, data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def validate_report(path: Path, key: str) -> dict[str, Any]:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    text = "\n".join([*paragraphs, *cells])
    required = [*REQUIRED_KEYWORDS, SYSTEMS[key]["label"]]
    if key != "yeast_gem":
        required.append("rxn1937")
    if key == "integrated":
        required.extend(["126 records/0 exact", "0 ready/8 unsupported", "construct hypothetical"])
    missing = [keyword for keyword in required if keyword not in text]
    if missing:
        raise ValueError(f"{path.name} missing required keywords after python-docx reopen: {missing}")
    code_runs = sum(
        1 for table in document.tables for row in table.rows for cell in row.cells
        for paragraph in cell.paragraphs for run in paragraph.runs if run.font.name == "Consolas"
    )
    if code_runs == 0:
        raise ValueError(f"{path.name} has no Consolas code runs")
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "table_rows": sum(len(table.rows) for table in document.tables),
        "inline_shapes": len(document.inline_shapes),
        "code_runs": code_runs,
        "keyword_check": "passed",
    }


def write_manifest(output_dir: Path, run_dir: Path, reports: dict[str, Path], data: dict[str, dict[str, Any]], stats: dict[str, dict[str, Any]]) -> Path:
    inputs: dict[str, str] = {}
    for system in data.values():
        for path in system["paths"].values():
            inputs[relative(path)] = sha256(path)
    source_paths = [
        ROOT / "08_runtime" / "run_10h2da_three_system_comparison.py",
        ROOT / "10_generic_target_workflow" / "runtime" / "workflow_cli.py",
        ROOT / "10_generic_target_workflow" / "runtime" / "executors.py",
    ]
    for path in source_paths:
        inputs[relative(path)] = sha256(path)
    manifest = {
        "manifest_version": "1.0",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "source_run": relative(run_dir),
        "generator": {"path": relative(Path(__file__)), "sha256": sha256(Path(__file__))},
        "reports": {
            key: {"path": relative(path), "sha256": sha256(path), "validation": stats[key]}
            for key, path in reports.items()
        },
        "input_artifact_hashes": dict(sorted(inputs.items())),
        "claims": {
            "exact_evidence_count": 0,
            "sequences_fabricated": False,
            "experimental_validation_claimed": False,
            "page_count_reported": False,
        },
    }
    path = output_dir / "10H2DA_three_system_report_manifest.json"
    path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return path


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate three reproducible Chinese 10H2DA Word reports from an existing comparison run.")
    parser.add_argument("--run-dir", type=Path, default=DEFAULT_RUN)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    run_dir = args.run_dir.resolve()
    output_dir = args.output_dir.resolve()
    if not run_dir.is_dir():
        raise SystemExit(f"comparison run directory does not exist: {run_dir}")
    data = {key: load_system(run_dir, key) for key in SYSTEMS}
    reports = {key: output_dir / info["filename"] for key, info in SYSTEMS.items()}
    for key, path in reports.items():
        build_report(data[key], run_dir, path)
    stats = {key: validate_report(path, key) for key, path in reports.items()}
    manifest = write_manifest(output_dir, run_dir, reports, data, stats)
    print(json.dumps({
        "status": "completed",
        "reports": {key: {"path": str(path), **stats[key]} for key, path in reports.items()},
        "manifest": str(manifest),
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
