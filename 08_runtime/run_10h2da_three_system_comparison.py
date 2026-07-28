from __future__ import annotations

import argparse
import json
import subprocess
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
CLI = ROOT / "10_generic_target_workflow" / "runtime" / "workflow_cli.py"
REFERENCE_CONFIG = ROOT / "10_generic_target_workflow" / "examples" / "target_workflow_10h2da_reference.json"

SYSTEMS = {
    "yeast_gem": {
        "label": "yeast-GEM",
        "role": "基线代谢模型",
        "model_path": "../Yeast-MetaTwin/Data/model/yeast-GEM.yml",
    },
    "yeast_metatwin": {
        "label": "Yeast-MetaTwin",
        "role": "地下代谢扩展模型",
        "model_path": "../Yeast-MetaTwin/Data/model/Yeast-MetaTwin.yml",
    },
    "integrated": {
        "label": "Integrated Deployment",
        "role": "综合预测部署系统（底层 GEM 为 Yeast-MetaTwin）",
        "model_path": "../Yeast-MetaTwin/Data/model/Yeast-MetaTwin.yml",
    },
}


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def make_configs(comparison_dir: Path) -> dict[str, Path]:
    reference = json.loads(REFERENCE_CONFIG.read_text(encoding="utf-8"))
    configs: dict[str, Path] = {}
    for key, system in SYSTEMS.items():
        config = json.loads(json.dumps(reference))
        config["target"]["target_id"] = f"10h2da_{key}"
        config["target"]["target_name"] = f"10H2DA comparison - {system['label']}"
        config["fba_validation"]["model_path"] = system["model_path"]
        config["comparison_context"] = {
            "system_id": key,
            "system_label": system["label"],
            "system_role": system["role"],
        }
        # comparison_context is provenance-only and not part of the strict runtime contract.
        config.pop("comparison_context")
        path = comparison_dir / "configs" / f"{key}.json"
        write_json(path, config)
        configs[key] = path
    return configs


def run_system(key: str, config: Path, comparison_dir: Path) -> dict[str, Any]:
    run_id = f"10h2da-{key}"
    runs_dir = comparison_dir / "runs"
    workspace = comparison_dir / "workspaces" / key
    command = [
        sys.executable,
        str(CLI),
        "run",
        "--config",
        str(config),
        "--project-root",
        str(ROOT),
        "--run-id",
        run_id,
        "--runs-dir",
        str(runs_dir),
        "--output-dir",
        str(workspace),
    ]
    completed = subprocess.run(command, cwd=ROOT, capture_output=True, text=True, encoding="utf-8", errors="replace")
    run_dir = runs_dir / run_id
    return {
        "key": key,
        "command": command,
        "returncode": completed.returncode,
        "stdout": completed.stdout.strip(),
        "stderr": completed.stderr.strip(),
        "run_dir": run_dir,
    }


def load_artifacts(run: dict[str, Any]) -> dict[str, Any]:
    run_dir: Path = run["run_dir"]
    required = {
        "state": run_dir / "state.json",
        "fba": run_dir / "artifacts" / "model_feasibility" / "fba_results.json",
    }
    for label, path in required.items():
        if not path.is_file():
            raise RuntimeError(f"{run['key']} missing {label} artifact: {path}\n{run['stderr']}")
    optional = {
        "kinetic": run_dir / "artifacts" / "kinetic_prediction" / "kinetic_predictions.json",
        "evidence": run_dir / "artifacts" / "external_evidence" / "external_evidence.json",
        "engineering": run_dir / "artifacts" / "engineering_feasibility" / "engineering_feasibility.json",
        "construct": run_dir / "artifacts" / "construct_design" / "construct_design.json",
    }
    artifacts = {name: json.loads(path.read_text(encoding="utf-8")) for name, path in required.items()}
    artifacts.update({name: json.loads(path.read_text(encoding="utf-8")) if path.is_file() else None for name, path in optional.items()})
    return artifacts


def route_row(system_key: str, route: dict[str, Any]) -> dict[str, Any]:
    fva_target = route.get("fva", {}).get("reactions", {}).get(route.get("objective_reaction_id"), {})
    sensitivity = {row["scenario"]: row for row in route.get("sensitivity", [])}
    return {
        "system": SYSTEMS[system_key]["label"],
        "route_id": route["route_id"],
        "status": route["status"],
        "objective": route.get("objective_value", 0.0),
        "fva_min": fva_target.get("minimum"),
        "fva_max": fva_target.get("maximum"),
        "pfba_total": route.get("pfba", {}).get("total_absolute_flux"),
        "yield": route.get("carbon_yield", {}).get("target_carbon_mol_per_substrate_carbon_mol"),
        "half_carbon": sensitivity.get("half_carbon", {}).get("objective_value"),
        "no_carbon": sensitivity.get("no_carbon", {}).get("objective_value"),
        "anaerobic": sensitivity.get("anaerobic", {}).get("objective_value"),
        "oxygen_bypass": sensitivity.get("anaerobic", {}).get("oxygen_bypass_suspected", False),
        "cycle": route.get("cycle_detection", {}).get("cycle_suspected", False),
        "stoichiometry": route.get("validation", {}).get("stoichiometric_feasibility", {}).get("status"),
        "enzymatic": route.get("validation", {}).get("enzymatic_validation", {}).get("status"),
    }


def fmt(value: Any, digits: int = 6) -> str:
    if value is None:
        return "N/A"
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def set_cell(cell, text: Any) -> None:
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(8.5)


def add_table(document: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell(cells[index], value)


def build_report(results: dict[str, dict[str, Any]], comparison_dir: Path, output_path: Path) -> dict[str, Any]:
    route_rows = [route_row(key, route) for key, artifacts in results.items() for route in artifacts["fba"]["routes"]]
    summary = {
        "generated_at": datetime.now().astimezone().isoformat(),
        "neutral_target_formula": "C10H18O3",
        "model_species_formula": "C10H17O3",
        "model_species_charge": -1,
        "comparison_dir": str(comparison_dir),
        "systems": {
            key: {
                "label": SYSTEMS[key]["label"],
                "role": SYSTEMS[key]["role"],
                "workflow_status": artifacts["state"]["status"],
                "model": artifacts["fba"]["model"],
                "native_growth_status": artifacts["fba"]["native_growth_status"],
                "native_growth_max": artifacts["fba"]["native_growth_max"],
                "routes": [row for row in route_rows if row["system"] == SYSTEMS[key]["label"]],
            }
            for key, artifacts in results.items()
        },
    }
    write_json(comparison_dir / "comparison_summary.json", summary)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10)

    title = document.add_heading("10-Hydroxy-trans-2-decenoic acid 三系统代谢通路预测比较报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("目标中性分子式：C10H18O3；模型物种：C10H17O3，电荷 -1")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("1. 执行摘要", level=1)
    document.add_paragraph(
        "本报告在统一目标定义、候选反应、培养条件、GLPK 求解器和 10% 最大生长下限下，"
        "并行运行 yeast-GEM、Yeast-MetaTwin 及 Integrated Deployment。前两项是代谢网络模型；"
        "第三项是综合预测部署系统，当前使用 Yeast-MetaTwin 作为底层 GEM，并额外执行证据、工程可行性和构建设计阶段。"
    )

    document.add_heading("2. 三个运行对象", level=1)
    add_table(
        document,
        ["运行对象", "定位", "实际底层模型", "工作流状态", "原生最大生长"],
        [
            [SYSTEMS[key]["label"], SYSTEMS[key]["role"], artifacts["fba"]["model"], artifacts["state"]["status"], fmt(artifacts["fba"]["native_growth_max"])]
            for key, artifacts in results.items()
        ],
    )

    document.add_heading("3. 候选通路", level=1)
    document.add_paragraph("路线 A（free_acid_route）：trans-dec-2-enoyl-CoA 水解为 trans-2-decenoate，再经 omega-羟基化形成 10H2DA。")
    document.add_paragraph("路线 B（coa_bound_route）：先对 trans-dec-2-enoyl-CoA 进行 omega-羟基化，再水解 10H2DA-CoA。")
    document.add_paragraph("四条候选反应在当前模型质子化约定下均通过元素和电荷平衡，但仍属于假设反应。")

    document.add_heading("4. FBA、pFBA 与 FVA 比较", level=1)
    add_table(
        document,
        ["系统", "路线", "状态", "最大目标通量", "90% FVA 下限", "90% FVA 上限", "pFBA 总通量", "C-mol 收率"],
        [
            [row["system"], row["route_id"], row["status"], fmt(row["objective"]), fmt(row["fva_min"]), fmt(row["fva_max"]), fmt(row["pfba_total"]), fmt(row["yield"])]
            for row in route_rows
        ],
    )

    document.add_heading("5. 培养条件与循环检查", level=1)
    add_table(
        document,
        ["系统", "路线", "半碳源通量", "无碳源通量", "禁氧交换通量", "氧旁路警告", "闭边界循环"],
        [
            [row["system"], row["route_id"], fmt(row["half_carbon"]), fmt(row["no_carbon"]), fmt(row["anaerobic"]), fmt(row["oxygen_bypass"]), fmt(row["cycle"])]
            for row in route_rows
        ],
    )

    gem_rows = [row for row in route_rows if row["system"] == "yeast-GEM"]
    meta_rows = [row for row in route_rows if row["system"] == "Yeast-MetaTwin"]
    document.add_heading("6. 模型差异解释", level=1)
    if gem_rows and meta_rows:
        for route_id in ("free_acid_route", "coa_bound_route"):
            gem = next(row for row in gem_rows if row["route_id"] == route_id)
            meta = next(row for row in meta_rows if row["route_id"] == route_id)
            delta = meta["objective"] - gem["objective"]
            document.add_paragraph(f"{route_id}：Yeast-MetaTwin 相对 yeast-GEM 的最大目标通量差为 {delta:.6f}。")
    document.add_paragraph(
        "Integrated Deployment 与 Yeast-MetaTwin 的 FBA 数值应一致，因为二者使用同一底层 YAML；"
        "Integrated 的新增价值体现在统一证据采集、插件门禁、工程规则、运行 provenance 和假设构建设计，而不是第三套独立计量网络。"
    )

    document.add_heading("7. 证据与可信度边界", level=1)
    integrated = results["integrated"]
    evidence = integrated.get("evidence") or {}
    kinetic = integrated.get("kinetic") or {}
    document.add_paragraph(
        f"综合系统本次证据阶段状态为 {evidence.get('status', 'N/A')}；"
        f"动力学阶段状态为 {kinetic.get('status', 'N/A')}。"
    )
    document.add_paragraph("FBA/FVA 只证明候选反应加入后网络在计量约束下可承载通量，不证明酶真实存在、具有底物特异性或能达到预测产量。")
    document.add_paragraph("禁用氧交换后仍出现目标通量时已标记 oxygen_bypass_suspected；该结果不能解释为真实厌氧可生产。")
    document.add_paragraph("候选反应当前无 GPR，单基因敲除分析为 not_applicable；需要精确 omega-hydroxylase 与 thioesterase 实验证据。")

    document.add_heading("8. 结论", level=1)
    best = max(route_rows, key=lambda row: row["objective"])
    document.add_paragraph(
        f"三套运行均完成。最高计算目标通量来自 {best['system']} 的 {best['route_id']}，数值为 {best['objective']:.6f}。"
        "最终路线优先级不能只按 FBA 通量决定，还必须结合酶证据、动力学适用域、氧旁路检查和实验验证。"
    )

    document.add_heading("9. 产物位置", level=1)
    document.add_paragraph(f"比较运行目录：{comparison_dir}")
    document.add_paragraph(f"机器可读摘要：{comparison_dir / 'comparison_summary.json'}")
    document.add_paragraph(f"Word 报告：{output_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return summary


def main() -> int:
    parser = argparse.ArgumentParser(description="Run and report the three-system 10H2DA comparison.")
    parser.add_argument("--comparison-id", default=datetime.now().strftime("10h2da-three-system-%Y%m%d-%H%M%S"))
    parser.add_argument("--output", type=Path, default=ROOT / "07_reports" / "10H2DA_three_system_comparison.docx")
    args = parser.parse_args()

    comparison_dir = ROOT / "runs" / args.comparison_id
    if comparison_dir.exists():
        raise SystemExit(f"comparison directory already exists: {comparison_dir}")
    configs = make_configs(comparison_dir)
    runs: dict[str, dict[str, Any]] = {}
    with ThreadPoolExecutor(max_workers=3) as pool:
        futures = {pool.submit(run_system, key, config, comparison_dir): key for key, config in configs.items()}
        for future in as_completed(futures):
            result = future.result()
            runs[result["key"]] = result

    failures = [result for result in runs.values() if result["returncode"] != 0]
    write_json(comparison_dir / "subprocess_results.json", {key: {**result, "run_dir": str(result["run_dir"])} for key, result in runs.items()})
    if failures:
        details = "\n".join(f"{item['key']}: exit {item['returncode']} {item['stderr']}" for item in failures)
        raise SystemExit(details)

    artifacts = {key: load_artifacts(runs[key]) for key in SYSTEMS}
    summary = build_report(artifacts, comparison_dir, args.output.resolve())
    print(json.dumps({"status": "completed", "comparison_dir": str(comparison_dir), "report": str(args.output.resolve()), "systems": list(summary["systems"])}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
