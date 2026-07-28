from __future__ import annotations

import argparse
import hashlib
import json
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import cobra
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RUN = ROOT / "runs" / "10h2da-three-system-20260717"
REPORT_DIR = ROOT / "07_reports"
OUTPUT = REPORT_DIR / "酿酒酵母_10H2DA_代谢预测综合报告.docx"
MANIFEST = REPORT_DIR / "酿酒酵母_10H2DA_代谢预测综合报告_manifest.json"

SYSTEMS = {
    "yeast_gem": ("yeast-GEM", "10h2da-yeast_gem"),
    "yeast_metatwin": ("Yeast-MetaTwin", "10h2da-yeast_metatwin"),
    "integrated": ("Integrated Deployment", "10h2da-integrated"),
}

EVALUATION_INPUTS = [
    "10h2da_terminal_evidence_validation.json",
    "10h2da_terminal_candidate_scores.csv",
    "10h2da_engineering_candidate_matrix.csv",
    "10h2da_engineering_candidate_prioritization_manifest.json",
    "10h2da_p450_engineering_feasibility_matrix.csv",
    "10h2da_p450_engineering_feasibility_manifest.json",
    "10h2da_external_evidence_supplement.json",
]

EXISTING_REPORTS = [
    "10H2DA_yeast-GEM_中文报告.docx",
    "10H2DA_Yeast-MetaTwin_中文报告.docx",
    "10H2DA_Integrated_Deployment_中文报告.docx",
]

REQUIRED_KEYWORDS = [
    "FAA2", "POX1", "TES1", "cyp52M1", "NCP1", "CYP505", "C10H18O3",
    "exact evidence count=0", "12项短板", "可复现代码", "r_0399", "r_0120",
    "r_2295", "r_2248", "rxn1937", "ACC1", "FAS1", "FAS2", "oxygen bypass",
]


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def rel(path: Path) -> str:
    try:
        return path.resolve().relative_to(ROOT.resolve()).as_posix()
    except ValueError:
        return str(path.resolve())


def set_font(run, name: str = "Microsoft YaHei", size: float = 9.5, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        properties.append(element)
    element.set(qn("w:fill"), fill)


def set_cell(cell, value: Any, header: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(value))
    set_font(run, size=7.5, bold=header, color="FFFFFF" if header else "24333A")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        shade(cell, "174E5C")


def add_table(document: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header, True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell(cells[index], value)
            if row_index % 2:
                shade(cells[index], "EAF1F2")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(document: Document, code: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "F1F3F4")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(code.splitlines()):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_font(run, name="Consolas", size=7.2, color="1D2529")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color in (
        ("Title", 21, "123D48"), ("Heading 1", 15, "123D48"), ("Heading 2", 11.5, "A8462C")
    ):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("10H2DA | 酿酒酵母代谢预测 | 综合报告"), size=8, color="55737B")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("来源：10h2da-three-system-20260717 | 计算结论须受证据等级约束"), size=8, color="55737B")


def collect_inputs() -> dict[str, Path]:
    paths: dict[str, Path] = {
        "comparison_summary": RUN / "comparison_summary.json",
        "subprocess_results": RUN / "subprocess_results.json",
        "comparison_script": ROOT / "08_runtime" / "run_10h2da_three_system_comparison.py",
        "workflow_cli": ROOT / "10_generic_target_workflow" / "runtime" / "workflow_cli.py",
        "executors": ROOT / "10_generic_target_workflow" / "runtime" / "executors.py",
    }
    for key, (_, run_id) in SYSTEMS.items():
        paths[f"config_{key}"] = RUN / "configs" / f"{key}.json"
        paths[f"fba_{key}"] = RUN / "runs" / run_id / "artifacts" / "model_feasibility" / "fba_results.json"
    for name in EVALUATION_INPUTS:
        paths[f"evaluation_{name}"] = ROOT / "06_evaluation" / name
    for name in EXISTING_REPORTS:
        paths[f"existing_report_{name}"] = REPORT_DIR / name
    missing = [str(path) for path in paths.values() if not path.is_file()]
    if missing:
        raise FileNotFoundError("Missing required inputs: " + ", ".join(missing))
    return paths


def load_models(configs: dict[str, dict[str, Any]]) -> tuple[dict[str, Any], list[dict[str, str]]]:
    models: dict[str, Any] = {}
    reaction_rows: list[dict[str, str]] = []
    for key in ("yeast_gem", "yeast_metatwin"):
        configured = Path(configs[key]["fba_validation"]["model_path"])
        model_path = configured if configured.is_absolute() else (ROOT / configured).resolve()
        model = cobra.io.load_yaml_model(model_path)
        models[key] = {
            "path": model_path,
            "model": model,
            "metabolites": len(model.metabolites),
            "reactions": len(model.reactions),
            "genes": len(model.genes),
        }
    for reaction_id in ("r_0399", "r_0120", "r_2295", "r_2248", "rxn1937"):
        model_key = "yeast_metatwin" if reaction_id == "rxn1937" else "yeast_gem"
        reaction = models[model_key]["model"].reactions.get_by_id(reaction_id)
        reaction_rows.append({
            "id": reaction.id,
            "equation": reaction.reaction,
            "gpr": reaction.gene_reaction_rule or "无 GPR",
            "genes": ", ".join(sorted(gene.id for gene in reaction.genes)),
        })
    return models, reaction_rows


def validate_existing_reports(paths: dict[str, Path]) -> list[list[str]]:
    expectations = {
        EXISTING_REPORTS[0]: ("yeast-GEM", "0.243015"),
        EXISTING_REPORTS[1]: ("Yeast-MetaTwin", "0.274011"),
        EXISTING_REPORTS[2]: ("Integrated Deployment", "0.274011"),
    }
    rows = []
    for name, (label, value) in expectations.items():
        path = paths[f"existing_report_{name}"]
        document = Document(path)
        text = "\n".join([p.text for p in document.paragraphs] + [c.text for t in document.tables for r in t.rows for c in r.cells])
        if label not in text or value not in text:
            raise ValueError(f"Existing report failed content check: {name}")
        rows.append([name, label, value, sha256(path), "已重新打开并核对"])
    return rows


def sensitivity(route: dict[str, Any], scenario: str) -> dict[str, Any]:
    return next(row for row in route["sensitivity"] if row["scenario"] == scenario)


def fmt(value: Any, digits: int = 6) -> str:
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def build_report(paths: dict[str, Path], output: Path) -> dict[str, Any]:
    summary = read_json(paths["comparison_summary"])
    configs = {key: read_json(paths[f"config_{key}"]) for key in SYSTEMS}
    fbas = {key: read_json(paths[f"fba_{key}"]) for key in SYSTEMS}
    terminal = read_json(paths["evaluation_10h2da_terminal_evidence_validation.json"])
    external = read_json(paths["evaluation_10h2da_external_evidence_supplement.json"])
    p450_manifest = read_json(paths["evaluation_10h2da_p450_engineering_feasibility_manifest.json"])
    engineering_manifest = read_json(paths["evaluation_10h2da_engineering_candidate_prioritization_manifest.json"])
    models, reaction_rows = load_models(configs)
    paths["model_yeast_gem"] = models["yeast_gem"]["path"]
    paths["model_yeast_metatwin"] = models["yeast_metatwin"]["path"]
    existing_report_rows = validate_existing_reports(paths)
    config = configs["yeast_metatwin"]

    document = Document()
    configure(document)
    title = document.add_heading("酿酒酵母 10H2DA 代谢预测综合报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Saccharomyces cerevisiae 生成 10-Hydroxy-trans-2-decenoic acid 的可追溯计算评估")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_font(run, size=11, color="A8462C")
    add_table(document, ["项目", "内容"], [
        ["目标中性分子", "10-Hydroxy-trans-2-decenoic acid (10H2DA), C10H18O3"],
        ["模型物种", "C10H17O3(-1)，单去质子化羧酸"],
        ["宿主", "Saccharomyces cerevisiae（酿酒酵母）"],
        ["来源运行", rel(RUN)],
        ["运行日期", summary["generated_at"]],
        ["求解器", "GLPK（config: glpk；FBA artifact executor: cobra_fba）"],
        ["报告生成时间", datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")],
    ])
    warning = document.add_paragraph("核心边界：终端反应 exact evidence count=0；本报告不把计算通量、家族相似性或工程评分表述为酶活或湿实验验证。")
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in warning.runs:
        set_font(run, bold=True, color="A8462C")

    document.add_section(WD_SECTION.NEW_PAGE)
    heading(document, "1. 执行摘要与结论等级")
    document.add_paragraph(
        "统一工作流在 yeast-GEM 与 Yeast-MetaTwin 上均得到正的候选 10H2DA 目标通量，且无碳源与闭边界检查未显示无物质来源的目标生成。"
        "Yeast-MetaTwin 的 0.274011 高于 yeast-GEM 的 0.243015，增幅约 12.75%，但其禁氧结果触发 oxygen bypass，不能据此宣称厌氧生产。"
        "两条终端路线的计量结果相同，酶学上优先游离酸路线；CoA-bound 羟化的底物证据更弱。"
    )
    add_table(document, ["结论层级", "状态", "可声明内容", "不可外推内容"], [
        ["软件可运行", "已通过", "三套 workflow 均 completed，产物可解析", "不代表模型或生物学正确"],
        ["计量可行", "已支持", "四条候选反应质量/电荷平衡，约束下有正通量", "不代表酶存在或有活性"],
        ["模型可行", "已支持但有条件", "10% 原生最大生长下两条路线可承载目标通量", "不代表真实产率；MetaTwin 有氧旁路"],
        ["酶学验证", "未验证", "仅到候选家族、近底物和工程优先级", "exact evidence count=0"],
        ["湿实验验证", "未验证", "尚无宿主生产、酶活或产物鉴定数据", "不得称已实现生物合成"],
    ])

    heading(document, "2. 目标化合物定义")
    add_table(document, ["对象", "分子式/电荷", "SMILES", "区室与身份假设"], [
        ["10H2DA 中性分子", "C10H18O3 / 0", "OCCCCCCC/C=C/C(=O)O", "化学身份；PubChem CID 5312738 isomeric SMILES"],
        ["模型 10H2DA", "C10H17O3 / -1", "OCCCCCCC/C=C/C(=O)O", "cand_10h2da_p；p 代表过氧化物酶体候选区室"],
        ["trans-2-decenoate", "C10H17O2 / -1", "CCCCCCC/C=C/C(=O)O", "cand_t2dec_p；模型质子约定"],
        ["trans-dec-2-enoyl-CoA", "C31H48N7O17P3S / -4", "见配置中的 PubChem CID 24883423 SMILES", "s_1507，过氧化物酶体"],
        ["10H2DA-CoA", "C31H48N7O18P3S / -4", "由 s_1507 作 omega 羟基替换推导", "内部计量定义，无直接 PubChem 命中"],
    ])
    document.add_paragraph("区室 p 是当前候选建模假设，不是转运或定位实验证据；游离酸、CoA 酯和 ER 膜 P450 之间的可达性必须单独验证。")

    heading(document, "3. 模型基础与角色")
    add_table(document, ["对象", "代谢物", "反应", "基因", "角色", "独立性"], [
        ["yeast-GEM", models["yeast_gem"]["metabolites"], models["yeast_gem"]["reactions"], models["yeast_gem"]["genes"], "基线、较保守的酵母 GEM", "独立底模 1"],
        ["Yeast-MetaTwin", models["yeast_metatwin"]["metabolites"], models["yeast_metatwin"]["reactions"], models["yeast_metatwin"]["genes"], "含地下/预测反应扩展的 GEM", "独立底模 2"],
        ["Integrated Deployment", models["yeast_metatwin"]["metabolites"], models["yeast_metatwin"]["reactions"], models["yeast_metatwin"]["genes"], "在 MetaTwin 上叠加证据、动力学门禁、工程规则和构建设计", "不是第三个独立 GEM"],
    ])
    document.add_paragraph("因此三系统是三次部署运行，但只有两套独立计量网络；Integrated 与 MetaTwin 同底模同 FBA，不能作为第三模型交叉验证。")

    heading(document, "4. 原生前体路径与模型节点")
    document.add_paragraph(
        "逐步路径：碳源进入中心碳代谢 -> acetyl-CoA/malonyl-CoA 与脂肪酸合成形成 C10 前体 -> decanoate -> "
        "decanoyl-CoA（r_0399，FAA2/YER015W）-> trans-dec-2-enoyl-CoA（r_0120，POX1/YGL205W）-> 候选终端反应 -> 10H2DA。"
        "r_2295（ECI1/YLR284C）提供替代连接；r_2248（FOX2/YKR009C）继续水合并竞争消耗该前体。"
    )
    role_map = {
        "r_0399": "C10 CoA 活化；模型支持",
        "r_0120": "trans-2-enoyl-CoA 形成；模型支持",
        "r_2295": "替代连接；模型支持，不是终端反应",
        "r_2248": "竞争消耗；模型支持，不自动推出敲除",
        "rxn1937": "MetaTwin 预测连接；完整 OR GPR，不是 14 基因同时必需",
    }
    add_table(document, ["反应", "模型方程", "实际完整 GPR", "基因", "判读"], [
        [row["id"], row["equation"], row["gpr"], row["genes"], role_map[row["id"]]] for row in reaction_rows
    ])
    document.add_paragraph(
        "ACC1、FAS1、FAS2 只能列为脂肪酸/C10 前体供给的工程候选；本次 FBA 未执行能证明其为 10H2DA 必需基因的因果分析，候选终端反应也无 GPR，单基因删除结果为 not_applicable。"
    )

    heading(document, "5. 两条候选终端路线")
    reaction_by_id = {row["reaction_id"]: row for row in config["candidate_reactions"]}
    terminal_meta = {
        "CAND_T2DEC_THIOESTERASE_P": ["游离酸 1", "TES1/PTE1 / YJR019C / P41903", "S. cerevisiae", "EC 3.1.2.-; 3.1.2.2", "H2O", "C：近底物/酶家族"],
        "CAND_T2DEC_OMEGA_HYDROXYLASE_P": ["游离酸 2", "cyp52M1 / B8QHP1 + NCP1/YHR042W/P16603；替代 CYP505/Q9Y8G7", "Starmerella bombicola；S. cerevisiae CPR；Fusarium oxysporum", "EC 1.14.-.-; 1.14.14.-；B8QHP1 1.14.14.80", "NADPH、O2、H+、heme；CYP52 需 CPR", "D：酶家族候选"],
        "CAND_T2DEC_COA_OMEGA_HYDROXYLASE_P": ["CoA-bound 1", "cyp52M1 / B8QHP1 + NCP1/YHR042W/P16603；替代 CYP505/Q9Y8G7", "同上", "EC 1.14.-.-; 1.14.14.-", "NADPH、O2、H+、heme；CYP52 需 CPR", "D：酶家族候选，风险更高"],
        "CAND_10H2DA_COA_THIOESTERASE_P": ["CoA-bound 2", "TES1/PTE1 / YJR019C / P41903", "S. cerevisiae", "EC 3.1.2.-; 3.1.2.2", "H2O", "D：酶家族候选"],
    }
    terminal_rows = []
    for route in config["routes"]:
        for reaction_id in route["reaction_ids"]:
            step, candidate, source, ec, cofactors, evidence = terminal_meta[reaction_id]
            reaction = reaction_by_id[reaction_id]
            terminal_rows.append([step, reaction_id, reaction["equation"], ec, cofactors, candidate, source, "质量/电荷均平衡", evidence])
    add_table(document, ["步骤", "反应 ID", "完整方程", "EC", "辅因子", "候选酶/基因/UniProt", "来源", "平衡", "证据等级"], terminal_rows)
    document.add_paragraph(
        "路线排序：优先先由 TES1/PTE1 候选释放 trans-2-decenoate，再由 cyp52M1/B8QHP1 候选羟化并配套 NCP1/P16603 或兼容 CPR。"
        "CYP505/Q9Y8G7 是 self-sufficient P450/还原酶融合替代项，可降低 CPR 配对风险。所有候选 exact evidence count=0；CoA-bound 羟化因已知家族证据偏游离脂肪酸而风险更高。"
    )
    add_table(document, ["候选", "本地可靠编号", "本地证据能支持什么", "不能支持什么"], [
        ["TES1/PTE1", "UniProt P41903；RHEA:16645；PMID 10092594", "acyl-CoA thioesterase 类及已知长链底物", "trans-dec-2-enoyl-CoA 或 10H2DA-CoA exact 活性"],
        ["cyp52M1", "UniProt B8QHP1；RHEA:56748；PMID 19054129, 23516968, 24242247", "真菌 CYP52 脂肪酸 omega 羟化家族", "trans-2-decenoate exact 转化"],
        ["CYP505/P450foxy", "UniProt Q9Y8G7；PMID 8830036, 11985584", "自足式 P450/CPR 融合与脂肪酸羟化语境", "10H2DA exact 反应"],
        ["NCP1/CPR", "UniProt P16603；基因 YHR042W", "酵母 NADPH-P450 reductase 配套候选", "本身不是 omega 羟化酶"],
        ["前体节点", "RHEA:33627, RHEA:40179, RHEA:28354", "decanoate/CoA 与 (2E)-decenoyl-CoA 邻近反应", "候选终端两步"],
    ])

    heading(document, "6. 三系统定量比较")
    comparison_rows = []
    for key, (label, _) in SYSTEMS.items():
        for route in fbas[key]["routes"]:
            objective = route["objective_reaction_id"]
            fva = route["fva"]["reactions"][objective]
            comparison_rows.append([
                label, route["route_id"], fmt(route["objective_value"]),
                f"[{fmt(fva['minimum'])}, {fmt(fva['maximum'])}] @90%",
                fmt(route["pfba"]["total_absolute_flux"]),
                fmt(route["carbon_yield"]["target_carbon_mol_per_substrate_carbon_mol"]),
                fmt(sensitivity(route, "half_carbon")["objective_value"]),
                fmt(sensitivity(route, "no_carbon")["objective_value"]),
                fmt(sensitivity(route, "anaerobic")["objective_value"]),
                str(sensitivity(route, "anaerobic")["oxygen_bypass_suspected"]),
                f"flux={fmt(route['cycle_detection']['target_flux'])}; suspected={route['cycle_detection']['cycle_suspected']}",
            ])
    add_table(document, ["系统", "路线", "FBA", "FVA", "pFBA 总绝对通量", "碳收率", "半碳源", "无碳源", "禁氧", "氧旁路", "闭边界循环"], comparison_rows)
    increase = (summary["systems"]["yeast_metatwin"]["routes"][0]["objective"] / summary["systems"]["yeast_gem"]["routes"][0]["objective"] - 1) * 100
    document.add_paragraph(
        f"yeast-GEM 为 0.243015；Yeast-MetaTwin 为 0.274011，约高 {increase:.2f}%。Integrated 采用同一 Yeast-MetaTwin YAML、同一 GLPK 和同一 FBA 实现，故数值相同。"
        "MetaTwin/Integrated 在禁用氧交换后仍有 0.274011，已标记 oxygen bypass；这提示模型内部产氧或替代氧连接，而非真实厌氧生物学。"
    )

    heading(document, "7. 异源生产所需功能模块")
    add_table(document, ["模块", "候选内容", "模型支持", "必须补充的验证"], [
        ["前体供给", "中心碳流、acetyl-CoA/malonyl-CoA、C10 链长控制；ACC1/FAS1/FAS2 为工程候选", "网络可到 C10 节点", "C10 池与碳流分配"],
        ["CoA 活化", "FAA2/r_0399 或适配 acyl-CoA ligase", "直接 GPR 支持", "底物谱、区室与能量负担"],
        ["trans-2 前体", "POX1/r_0120；ECI1/r_2295 替代连接", "直接 GPR 支持", "前体积累与竞争分流"],
        ["终端酶", "TES1/PTE1 + cyp52M1；CYP505 替代", "候选反应计量可行", "exact 底物酶活与位置选择性"],
        ["P450 配套", "NCP1/兼容 CPR、heme、NADPH、O2", "辅因子计量进入反应", "电子耦联、表达与氧化压力"],
        ["区室运输", "过氧化物酶体、胞质、ER 间 CoA 酯/游离酸可达性", "当前仅 p 区室假设", "定位与转运机制"],
        ["产物输出/耐受", "10H2DA 输出、胞内积累与宿主耐受", "本次未建模", "毒性、输出与稳态"],
        ["检测验证", "区分 trans 异构体、游离酸/CoA 形式并核对质量与碳平衡", "无湿实验数据", "外部分析确认"],
    ])
    document.add_paragraph("本节只定义功能需求和验证问题，不提供培养、表达、剂量、取样或分析的逐步湿实验协议与参数。")

    heading(document, "8. 基因分层与设计优先级")
    add_table(document, ["层级", "对象", "优先级", "证据标签", "结论边界"], [
        ["模型直接基因", "FAA2/YER015W；POX1/YGL205W", "高", "模型支持", "前体路径节点，不等于必须改造"],
        ["模型替代/竞争", "ECI1/YLR284C；FOX2/YKR009C；rxn1937 ORF 集合", "按通量与区室复核", "模型支持", "替代连接/竞争节点，不自动推出敲除"],
        ["宿主终端候选", "TES1/PTE1/YJR019C/P41903", "优先", "候选待验证", "酶类和近底物支持，exact=0"],
        ["异源候选", "cyp52M1/B8QHP1", "优先", "候选待验证", "真菌 CYP52 工程评分高，非 exact 活性"],
        ["异源替代", "CYP505/Q9Y8G7", "次优先并行", "候选待验证", "self-sufficient，仍需宿主兼容和底物验证"],
        ["辅因子候选", "NCP1/YHR042W/P16603；heme/NADPH/O2 供给", "配套", "候选待验证", "支持 P450，不是终端羟化酶"],
        ["前体工程", "ACC1、FAS1、FAS2", "建模后排序", "候选待验证", "非本次 FBA 证明的必需基因"],
        ["低特异性不推荐", "ERG11、ERG5、DIT2", "低/不推荐", "关键词或大类命中", "原生底物类别距离远，不应当作 10H2DA 羟化证据"],
    ])

    heading(document, "9. 12项短板与关闭条件")
    gaps = [
        [1, "终端酶", "四个反应 exact evidence count=0", "扩大经人工审阅的底物-酶证据与序列候选", "是"],
        [2, "GPR", "候选终端反应无 GPR", "将验证后的序列映射到候选反应并做基因扰动模拟", "是"],
        [3, "区室", "p 区室是假设，ER P450 与过氧化物酶体底物可达性未证", "建立分区反应/转运并做可达性分析", "是"],
        [4, "氧旁路", "MetaTwin 禁氧仍有通量", "追踪产氧反应、收紧氧边界、做热力学/loopless 复核", "否；生物学确认仍需实验"],
        [5, "两路线排序", "FBA 通量相同不能区分酶学可行性", "按底物证据、区室、辅因子和风险加权排序", "是"],
        [6, "动力学", "缺 exact 底物 kcat/Km；预测适用域有限", "获取序列与可靠底物动力学后约束 ecModel/采样", "是"],
        [7, "热力学", "未做 ΔG 和浓度方向性约束", "补充组分贡献/浓度范围并做 thermodynamic FBA", "外部数据通常需要"],
        [8, "前体设计", "ACC1/FAS1/FAS2 尚未由本次 FBA 排序", "做产物耦联、通量扫描和基因扰动比较", "否；最终确认需实验"],
        [9, "P450 配套", "CPR/heme/NADPH/O2 与表达耦联未知", "分别约束电子、heme、氧和 NADPH 预算并比较 CYP52/CYP505", "是"],
        [10, "数据质量", "地下反应与候选标签含预测来源，质量不均", "逐条审计平衡、provenance、重复与证据等级", "否；酶学声明需实验"],
        [11, "外部验证", "无产物、酶活和宿主生产数据", "用独立分析确认化学身份、酶活和宿主表型", "是"],
        [12, "第三模型独立性", "Integrated 复用 MetaTwin，不是第三 GEM", "引入真正独立、版本固定且同条件的第三模型", "否；但模型结论仍需实验"],
    ]
    add_table(document, ["编号", "短板", "当前问题", "可计算/工程补救", "必须外部实验才能关闭"], gaps)

    heading(document, "10. 分阶段决策路径")
    add_table(document, ["阶段", "输入", "决策门", "输出"], [
        ["1 证据确认", "本地 UniProt/Rhea/PMID 与候选矩阵", "区分 exact、近底物、家族与关键词证据", "保留/淘汰候选及证据等级"],
        ["2 序列与动力学", "可追溯序列、底物身份与适用域", "无序列或域外预测不得升格", "可约束的 kcat/Km 范围与不确定性"],
        ["3 模型约束", "GPR、区室、运输、热力学、P450 辅因子", "关闭氧旁路并复核两路线", "鲁棒通量范围和设计假设"],
        ["4 菌株设计", "模型支持节点与候选待验证模块", "生长、前体、辅因子和竞争分流共同评估", "分层设计优先级，不含操作参数"],
        ["5 实验验证", "优先候选和明确判定标准", "化学身份、酶活、区室和宿主生产均需独立证据", "关闭或保留关键短板"],
    ])

    heading(document, "11. 证据层与既有报告复核")
    add_table(document, ["证据产物", "实际规模/结果", "本报告用法"], [
        ["terminal evidence", f"direct match={terminal['direct_10h2da_match_count']}；候选 verdict={len(terminal['candidate_verdicts'])}", "限定终端酶为 enzyme_class_support_only"],
        ["external supplement", f"records={external['record_count']}；UniProt/Rhea/PubMed={external['source_counts']}", "只引用本地存在的 accession、Rhea、PMID"],
        ["engineering prioritization", f"combined rows={engineering_manifest['combined_rows']}；design rows={engineering_manifest['design_rows']}", "作为工程排序，不作为酶学证明"],
        ["P450 feasibility", f"hydroxylase rows={p450_manifest['hydroxylase_rows']}；design rows={p450_manifest['design_rows']}", "比较 CPR、表达、底物与宿主风险"],
    ])
    add_table(document, ["既有独立报告", "系统标识", "核对数值", "SHA256", "读取状态"], existing_report_rows)
    document.add_paragraph("三份既有 Word 报告作为已读取的二级解释输入；所有关键数值仍以 comparison_summary.json 和三个 fba_results.json 为主记录。")

    heading(document, "12. 可复现代码与命令")
    document.add_paragraph("三系统比较脚本命令（会新建 run 并重跑 FBA；本次生成综合报告未执行该命令）：")
    add_code(document, f'''python "{(ROOT / '08_runtime' / 'run_10h2da_three_system_comparison.py').resolve()}" --comparison-id 10h2da-three-system-REPRODUCE --output "{(REPORT_DIR / '10H2DA_three_system_comparison_REPRODUCE.docx').resolve()}"''')
    document.add_paragraph("workflow CLI 准确命令（分别对应既有三系统配置；使用新 run-id/目录可避免覆盖）：")
    workflow_commands = []
    for key, (_, _) in SYSTEMS.items():
        workflow_commands.append(
            f'python "{(ROOT / "10_generic_target_workflow" / "runtime" / "workflow_cli.py").resolve()}" run '
            f'--config "{(RUN / "configs" / f"{key}.json").resolve()}" --project-root "{ROOT.resolve()}" '
            f'--run-id 10h2da-{key}-REPRODUCE --runs-dir "{(RUN / "report-reruns").resolve()}" '
            f'--output-dir "{(RUN / "report-rerun-workspaces" / key).resolve()}"'
        )
    add_code(document, "\n".join(workflow_commands))
    document.add_paragraph("COBRApy 模型查询代码：")
    add_code(document, f'''from pathlib import Path
import cobra

paths = [
    Path(r"{models['yeast_gem']['path']}"),
    Path(r"{models['yeast_metatwin']['path']}"),
]
for path in paths:
    model = cobra.io.load_yaml_model(path)
    print(path.name, len(model.metabolites), len(model.reactions), len(model.genes))
    ids = ["r_0399", "r_0120", "r_2295", "r_2248"]
    if path.name == "Yeast-MetaTwin.yml":
        ids.append("rxn1937")
    for reaction_id in ids:
        reaction = model.reactions.get_by_id(reaction_id)
        print(reaction.id, reaction.reaction, reaction.gene_reaction_rule)''')
    document.add_paragraph("结果解析代码：")
    add_code(document, f'''import json
from pathlib import Path

root = Path(r"{RUN.resolve()}")
summary = json.loads((root / "comparison_summary.json").read_text(encoding="utf-8"))
for key, system in summary["systems"].items():
    for route in system["routes"]:
        print(key, route["route_id"], route["objective"], route["fva_min"], route["fva_max"],
              route["pfba_total"], route["yield"], route["half_carbon"], route["no_carbon"],
              route["anaerobic"], route["oxygen_bypass"], route["cycle"])

for key, run_id in {repr({key: run_id for key, (_, run_id) in SYSTEMS.items()})}.items():
    path = root / "runs" / run_id / "artifacts" / "model_feasibility" / "fba_results.json"
    result = json.loads(path.read_text(encoding="utf-8"))
    print(key, result["native_growth_status"], result["native_growth_max"], result["routes"])''')
    document.add_paragraph("综合报告再生成命令（只读既有产物，不重跑 FBA）：")
    add_code(document, f'''python "{Path(__file__).resolve()}" --output "{output.resolve()}" --manifest "{MANIFEST.resolve()}"''')

    heading(document, "13. 路径、provenance 与限制")
    add_table(document, ["类别", "准确路径"], [
        ["比较摘要", str(paths["comparison_summary"].resolve())],
        ["yeast-GEM FBA", str(paths["fba_yeast_gem"].resolve())],
        ["MetaTwin FBA", str(paths["fba_yeast_metatwin"].resolve())],
        ["Integrated FBA", str(paths["fba_integrated"].resolve())],
        ["终端证据", str(paths["evaluation_10h2da_terminal_evidence_validation.json"].resolve())],
        ["工程候选", str(paths["evaluation_10h2da_engineering_candidate_matrix.csv"].resolve())],
        ["P450 feasibility", str(paths["evaluation_10h2da_p450_engineering_feasibility_matrix.csv"].resolve())],
        ["生成器", str(Path(__file__).resolve())],
    ])
    document.add_paragraph(
        "限制汇总：目标不存在于原生模型；终端反应是平衡的候选扩展；FVA 非 loopless；氧旁路尚未关闭；动力学预测不能替代 exact 底物数据；"
        "区室、运输、表达、辅因子耦联、产物输出、耐受和分析身份均未获湿实验确认；Integrated 不提供第三套独立 GEM。"
    )
    document.add_paragraph("所有输入 SHA256、生成器 SHA256、DOCX SHA256、运行日期与自检统计记录于同名 JSON manifest。")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return {
        "models": {key: {name: value for name, value in row.items() if name != "model"} for key, row in models.items()},
        "run_generated_at": summary["generated_at"],
        "solver": "glpk",
        "exact_evidence_count": terminal["direct_10h2da_match_count"],
    }


def validate_report(path: Path) -> dict[str, Any]:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    text = "\n".join([*paragraphs, *cells])
    missing = [keyword for keyword in REQUIRED_KEYWORDS if keyword not in text]
    if missing:
        raise ValueError(f"Missing required keywords after python-docx reopen: {missing}")
    if len(document.tables) < 10:
        raise ValueError(f"Expected at least 10 tables, found {len(document.tables)}")
    code_runs = sum(
        1 for table in document.tables for row in table.rows for cell in row.cells
        for paragraph in cell.paragraphs for run in paragraph.runs if run.font.name == "Consolas"
    )
    if code_runs == 0:
        raise ValueError("No Consolas code runs found")
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "table_rows": sum(len(table.rows) for table in document.tables),
        "inline_shapes": len(document.inline_shapes),
        "code_runs": code_runs,
        "keyword_check": "passed",
        "minimum_table_check": "passed",
        "consolas_check": "passed",
    }


def write_manifest(path: Path, output: Path, inputs: dict[str, Path], stats: dict[str, Any], provenance: dict[str, Any]) -> None:
    payload = {
        "manifest_version": "1.0",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "report": {"path": rel(output), "sha256": sha256(output), "validation": stats},
        "generator": {"path": rel(Path(__file__)), "sha256": sha256(Path(__file__))},
        "source_run": rel(RUN),
        "run_generated_at": provenance["run_generated_at"],
        "solver": provenance["solver"],
        "exact_evidence_count": provenance["exact_evidence_count"],
        "model_summary": {
            key: {name: rel(value) if name == "path" else value for name, value in row.items()}
            for key, row in provenance["models"].items()
        },
        "input_artifact_hashes": {
            rel(input_path): sha256(input_path) for input_path in sorted(set(inputs.values()), key=lambda item: rel(item))
        },
        "claims": {
            "stoichiometric_feasibility_supported": True,
            "enzymatic_validation_established": False,
            "wet_lab_validation_established": False,
            "integrated_is_independent_third_gem": False,
            "fba_rerun_for_report_generation": False,
        },
    }
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str) + "\n", encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate the standalone Chinese comprehensive 10H2DA report from existing artifacts.")
    parser.add_argument("--output", type=Path, default=OUTPUT)
    parser.add_argument("--manifest", type=Path, default=MANIFEST)
    args = parser.parse_args()
    inputs = collect_inputs()
    provenance = build_report(inputs, args.output.resolve())
    stats = validate_report(args.output.resolve())
    write_manifest(args.manifest.resolve(), args.output.resolve(), inputs, stats, provenance)
    print(json.dumps({
        "status": "completed",
        "report": str(args.output.resolve()),
        "manifest": str(args.manifest.resolve()),
        "report_sha256": sha256(args.output.resolve()),
        "validation": stats,
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
